from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List
import uuid
import os
from datetime import datetime
from docx import Document
from supabase import create_client, Client as SupabaseClient

from ..config import settings
from ..database import get_db
from ..models.offer import Offer, Contract, OfferStatus, ContractStatus
from ..models.client import Client
from ..models.user import User
from ..models.vehicle import Vehicle
from ..schemas.offer import OfferCreate, OfferResponse, ContractResponse, ContractCreateRequest
from ..api.auth import get_current_user

# --- AUTH BYPASS FOR LOCAL DEV (So the UI doesn't break due to missing JWT) ---
def mock_get_current_user(db: Session = Depends(get_db)):
    user = db.query(User).first()
    if not user:
        user = User(email="admin@axis.ro", hashed_password="mock", full_name="Mock Admin", role="Super Admin")
        db.add(user)
        db.commit()
        db.refresh(user)
    return user
# -----------------------------------------------------------------------------

router = APIRouter(prefix="/api/offers", tags=["Offers & Contracts"])

def calculate_monthly_rate(price, advance_pct, residual_pct, period, interest_rate):
    """
    Mock financial calculation.
    In real life this would use PMT formula.
    """
    advance = price * (advance_pct / 100)
    residual = price * (residual_pct / 100)
    financed_amount = price - advance - residual
    
    # Simple mock calculation (Principal + total interest) / period
    total_interest = financed_amount * (interest_rate / 100) * (period / 12)
    monthly_payment = (financed_amount + total_interest) / period
    return round(monthly_payment, 2)

@router.post("/", response_model=OfferResponse)
def create_offer(offer: OfferCreate, db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    # Check client
    client = db.query(Client).filter(Client.id == offer.client_id).first()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
        
    rate = calculate_monthly_rate(
        offer.vehicle_price, 
        offer.advance_percent, 
        offer.residual_value_percent, 
        offer.period_months, 
        offer.interest_rate
    )
    
    new_offer = Offer(
        **offer.model_dump(),
        monthly_rate=rate,
        created_by_id=current_user.id
    )
    db.add(new_offer)
    db.commit()
    db.refresh(new_offer)
    return new_offer

from sqlalchemy.orm import joinedload

@router.get("/", response_model=List[OfferResponse])
def get_offers(db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    offers = db.query(Offer).options(joinedload(Offer.contract)).order_by(Offer.created_at.desc()).all()
    return offers

@router.get("/{offer_id}", response_model=OfferResponse)
def get_offer(offer_id: int, db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    offer = db.query(Offer).options(joinedload(Offer.contract)).filter(Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
    return offer

@router.put("/{offer_id}", response_model=OfferResponse)
def update_offer(offer_id: int, request: OfferCreate, db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    offer = db.query(Offer).filter(Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
    
    if offer.status != OfferStatus.DRAFT:
        raise HTTPException(status_code=400, detail="Doar ofertele Draft pot fi editate.")
        
    offer.client_id = request.client_id
    offer.vehicle_id = getattr(request, 'vehicle_id', None)
    offer.vehicle_make = request.vehicle_make
    offer.vehicle_model = request.vehicle_model
    offer.vehicle_price = request.vehicle_price
    offer.advance_percent = request.advance_percent
    offer.period_months = request.period_months
    offer.residual_value_percent = request.residual_value_percent
    offer.interest_rate = request.interest_rate
    
    advance = (offer.vehicle_price * offer.advance_percent) / 100
    residual = (offer.vehicle_price * offer.residual_value_percent) / 100
    financed = offer.vehicle_price - advance - residual
    total_interest = financed * (offer.interest_rate / 100) * (offer.period_months / 12)
    offer.monthly_rate = (financed + total_interest) / offer.period_months

    db.commit()
    db.refresh(offer)
    return offer

@router.post("/{offer_id}/approve", response_model=OfferResponse)
def approve_offer(offer_id: int, db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    offer = db.query(Offer).filter(Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
        
    offer.status = OfferStatus.APPROVED
    offer.approved_by_id = current_user.id
    db.commit()
    db.refresh(offer)
    return offer

@router.post("/upload-template")
async def upload_template(file: UploadFile = File(...), current_user = Depends(mock_get_current_user)):
    os.makedirs("templates", exist_ok=True)
    file_location = f"templates/contract_template.docx"
    with open(file_location, "wb+") as file_object:
        file_object.write(file.file.read())
    return {"info": f"file '{file.filename}' saved at '{file_location}'"}

@router.post("/{offer_id}/generate-contract", response_model=ContractResponse)
def generate_contract(offer_id: int, request: ContractCreateRequest, db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    offer = db.query(Offer).filter(Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
    
    vehicle = None
    vehicle_id_to_use = request.vehicle_id or offer.vehicle_id
    if vehicle_id_to_use:
        vehicle = db.query(Vehicle).filter(Vehicle.id == vehicle_id_to_use).first()
        if not vehicle:
            raise HTTPException(status_code=404, detail="Vehicle not found")
        
    contract_num = f"AXIS-{datetime.utcnow().year}-{uuid.uuid4().hex[:6].upper()}"
    document_url = f"/documents/{contract_num}.docx"
    document_path = f"documents/{contract_num}.docx"
    
    template_path = "templates/contract_template.docx"
    
    if os.path.exists(template_path):
        doc = Document(template_path)
        replacements = {
            "{{client_name}}": offer.client.name,
            "{{client_cui}}": offer.client.cui_cnp,
            "{{client_address}}": offer.client.address or "",
            "{{vehicle_make}}": vehicle.make if vehicle else offer.vehicle_make,
            "{{vehicle_model}}": vehicle.model if vehicle else offer.vehicle_model,
            "{{vehicle_vin}}": vehicle.vin if vehicle else "___________",
            "{{vehicle_plate}}": vehicle.license_plate if vehicle else "___________",
            "{{vehicle_price}}": str(offer.vehicle_price),
            "{{monthly_rate}}": str(offer.monthly_rate),
            "{{period_months}}": str(offer.period_months),
            "{{nr_contract}}": contract_num
        }
        
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
                                
        os.makedirs("documents", exist_ok=True)
        doc.save(document_path)
        
        # Upload to Supabase Storage if configured
        if settings.SUPABASE_URL and settings.SUPABASE_SERVICE_KEY:
            try:
                supabase: SupabaseClient = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_KEY)
                file_name = f"{contract_num}.docx"
                
                with open(document_path, 'rb') as f:
                    supabase.storage.from_('axis-documents').upload(
                        file_name, 
                        f, 
                        {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                    )
                
                # Get the public URL from Supabase
                document_url = supabase.storage.from_('axis-documents').get_public_url(file_name)
                print(f"Contract uploaded to Supabase: {document_url}")
                
            except Exception as e:
                print(f"Supabase upload failed, falling back to local: {e}")
    else:
        doc = Document()
        doc.add_heading(f'Contract Auto - {contract_num}', 0)
        doc.add_paragraph(f'Client: {offer.client.name}')
        doc.add_paragraph(f'Vehicul: {vehicle.make if vehicle else offer.vehicle_make} {vehicle.model if vehicle else offer.vehicle_model}')
        doc.add_paragraph(f'Pret: {offer.vehicle_price} {offer.currency if hasattr(offer, "currency") else "EUR"}')
        output_path = f"documents/{contract_num}.docx"
        doc.save(output_path)
        document_url = f"/documents/{contract_num}.docx"
    
    new_contract = Contract(
        offer_id=offer.id,
        vehicle_id=vehicle_id_to_use,
        contract_number=contract_num,
        status=ContractStatus.GENERATED,
        document_url=document_url
    )
    offer.status = OfferStatus.CONVERTED
    
    db.add(new_contract)
    db.commit()
    db.refresh(new_contract)
    return new_contract

@router.get("/contracts", response_model=List[ContractResponse])
def get_contracts(db: Session = Depends(get_db), current_user = Depends(mock_get_current_user)):
    return db.query(Contract).order_by(Contract.created_at.desc()).all()
